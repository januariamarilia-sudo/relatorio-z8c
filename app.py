from __future__ import annotations

import copy
import json
import os
import re
import socket
import sys
import threading
import webbrowser
from datetime import datetime
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, urlparse

from docx import Document


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "modelo_relatorio.docx"
OUTPUT_DIR = BASE_DIR / "saida"
HOST = os.environ.get("RELATORIO_HOST", "127.0.0.1")
DEFAULT_PORT = 8765


HTML = r"""<!doctype html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Gerador de Relatório de Produtividade</title>
  <style>
    :root {
      --ink: #172033;
      --muted: #5c6576;
      --line: #d9dee8;
      --soft: #f4f6fa;
      --brand: #1868a8;
      --brand-dark: #0e4b7d;
      --ok: #16724a;
      --danger: #a83232;
      --white: #ffffff;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      font-family: Arial, Helvetica, sans-serif;
      color: var(--ink);
      background: #eef2f7;
      letter-spacing: 0;
    }
    header {
      position: sticky;
      top: 0;
      z-index: 10;
      background: var(--white);
      border-bottom: 1px solid var(--line);
    }
    .topbar {
      max-width: 1180px;
      margin: 0 auto;
      padding: 14px 18px;
      display: flex;
      gap: 12px;
      align-items: center;
      justify-content: space-between;
    }
    h1 {
      font-size: 20px;
      line-height: 1.2;
      margin: 0;
      font-weight: 700;
    }
    main {
      max-width: 1180px;
      margin: 0 auto;
      padding: 18px;
    }
    .actions {
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      justify-content: flex-end;
    }
    button, .button {
      border: 1px solid var(--line);
      background: var(--white);
      color: var(--ink);
      min-height: 38px;
      padding: 8px 12px;
      border-radius: 6px;
      font-weight: 700;
      cursor: pointer;
      text-decoration: none;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      gap: 7px;
    }
    button.primary {
      background: var(--brand);
      border-color: var(--brand);
      color: var(--white);
    }
    button.primary:hover { background: var(--brand-dark); }
    button.danger { color: var(--danger); }
    button.small {
      min-height: 32px;
      padding: 6px 9px;
      font-weight: 700;
    }
    section {
      background: var(--white);
      border: 1px solid var(--line);
      border-radius: 8px;
      margin-bottom: 14px;
      overflow: hidden;
    }
    .section-head {
      padding: 13px 14px;
      background: var(--soft);
      border-bottom: 1px solid var(--line);
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
    }
    h2 {
      font-size: 16px;
      margin: 0;
      line-height: 1.2;
    }
    .body { padding: 14px; }
    .grid {
      display: grid;
      grid-template-columns: repeat(4, minmax(0, 1fr));
      gap: 12px;
    }
    label {
      display: block;
      font-size: 12px;
      font-weight: 700;
      color: var(--muted);
      margin-bottom: 5px;
    }
    input, select, textarea {
      width: 100%;
      border: 1px solid var(--line);
      border-radius: 6px;
      padding: 9px 10px;
      font: inherit;
      color: var(--ink);
      background: var(--white);
    }
    textarea {
      min-height: 92px;
      resize: vertical;
      line-height: 1.35;
    }
    .wide { grid-column: span 2; }
    .full { grid-column: 1 / -1; }
    .table-wrap {
      overflow-x: auto;
      border: 1px solid var(--line);
      border-radius: 8px;
    }
    table {
      width: 100%;
      border-collapse: collapse;
      min-width: 880px;
      table-layout: fixed;
    }
    th, td {
      border-bottom: 1px solid var(--line);
      padding: 7px;
      vertical-align: top;
    }
    th {
      background: #f8fafc;
      text-align: left;
      font-size: 12px;
      color: var(--muted);
    }
    td input, td select, td textarea {
      border-color: transparent;
      background: transparent;
      padding: 6px;
      min-height: 34px;
    }
    td textarea { min-height: 44px; }
    tr:last-child td { border-bottom: 0; }
    .col-num { width: 54px; }
    .col-actions { width: 54px; text-align: center; }
    .status {
      min-height: 40px;
      display: flex;
      align-items: center;
      gap: 8px;
      color: var(--muted);
      font-size: 13px;
    }
    .status.ok { color: var(--ok); font-weight: 700; }
    .indicator-grid {
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 12px;
    }
    .indicator {
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 10px;
      background: #fbfcfe;
    }
    .indicator label { margin-bottom: 7px; }
    .two {
      display: grid;
      grid-template-columns: 180px minmax(0, 1fr);
      gap: 8px;
      align-items: start;
    }
    .download-line {
      margin-top: 10px;
      display: none;
      gap: 10px;
      align-items: center;
      flex-wrap: wrap;
    }
    .download-line.show { display: flex; }
    @media (max-width: 840px) {
      .topbar { align-items: stretch; flex-direction: column; }
      .actions { justify-content: flex-start; }
      .grid, .indicator-grid, .two { grid-template-columns: 1fr; }
      .wide { grid-column: 1 / -1; }
      main { padding: 12px; }
    }
  </style>
</head>
<body>
  <header>
    <div class="topbar">
      <h1>Gerador de Relatório de Produtividade</h1>
      <div class="actions">
        <button id="copyPlan" type="button">Copiar planejamento para executadas</button>
        <button id="generate" class="primary" type="button">Gerar Word</button>
      </div>
    </div>
  </header>
  <main>
    <section>
      <div class="section-head"><h2>Identificação</h2></div>
      <div class="body">
        <div class="grid">
          <div><label>Data ou período</label><input id="data" placeholder="11/07/2026"></div>
          <div><label>Colaborador(a)</label><input id="colaborador"></div>
          <div><label>Setor/Núcleo</label><input id="setor"></div>
          <div><label>Função</label><input id="funcao"></div>
          <div><label>Horário de trabalho</label><input id="horario"></div>
          <div><label>Responsável pela validação</label><input id="responsavel"></div>
        </div>
        <div id="result" class="download-line">
          <span id="resultText" class="status ok"></span>
          <a id="download" class="button" href="#">Baixar relatório</a>
        </div>
      </div>
    </section>

    <section>
      <div class="section-head">
        <h2>Planejamento semanal</h2>
        <button class="small" type="button" data-add="planning">Adicionar linha</button>
      </div>
      <div class="body"><div class="table-wrap"><table id="planning"></table></div></div>
    </section>

    <section>
      <div class="section-head">
        <h2>Demandas extraordinárias</h2>
        <button class="small" type="button" data-add="extra">Adicionar linha</button>
      </div>
      <div class="body"><div class="table-wrap"><table id="extra"></table></div></div>
    </section>

    <section>
      <div class="section-head">
        <h2>Atividades executadas</h2>
        <button class="small" type="button" data-add="done">Adicionar linha</button>
      </div>
      <div class="body"><div class="table-wrap"><table id="done"></table></div></div>
    </section>

    <section>
      <div class="section-head"><h2>Indicadores qualitativos do dia</h2></div>
      <div class="body indicator-grid">
        <div class="indicator"><label>Nível de produtividade percebido</label><select id="produtividade"><option>Alto</option><option>Médio</option><option>Baixo</option></select></div>
        <div class="indicator"><label>Cumprimento das atividades planejadas</label><select id="cumprimento"><option>Sim</option><option>Parcialmente</option><option>Não</option></select></div>
        <div class="indicator two"><div><label>Demanda urgente não planejada?</label><select id="urgente"><option>Sim</option><option>Não</option></select></div><div><label>Comentário</label><input id="urgente_comentario"></div></div>
        <div class="indicator two"><div><label>Dependência de outro setor?</label><select id="dependencia"><option>Sim</option><option>Não</option></select></div><div><label>Qual?</label><input id="dependencia_comentario"></div></div>
        <div class="indicator two"><div><label>Houve retrabalho?</label><select id="retrabalho"><option>Sim</option><option>Não</option></select></div><div><label>Por quê?</label><input id="retrabalho_comentario"></div></div>
        <div class="indicator two"><div><label>Houve sobrecarga?</label><select id="sobrecarga"><option>Sim</option><option>Não</option></select></div><div><label>Discorra</label><input id="sobrecarga_comentario"></div></div>
        <div class="indicator two"><div><label>Houve acúmulo de demandas?</label><select id="acumulo"><option>Sim</option><option>Não</option></select></div><div><label>Discorra</label><input id="acumulo_comentario"></div></div>
      </div>
    </section>

    <section>
      <div class="section-head"><h2>Observações e validação</h2></div>
      <div class="body grid">
        <div class="wide"><label>Observações do(a) colaborador(a)</label><textarea id="observacoes"></textarea></div>
        <div class="wide"><label>Análise/validação da chefia imediata</label><textarea id="validacao"></textarea></div>
      </div>
    </section>
  </main>

  <script>
    const schemas = {
      planning: [
        ["processo", "Processo/assunto", "textarea"],
        ["descricao", "Descrição resumida", "textarea"],
        ["prioridade", "Prioridade", "select", ["Alta", "Média", "Baixa", "Urgente"]],
        ["status", "Status", "input"]
      ],
      extra: [
        ["processo", "Processo/assunto", "textarea"],
        ["descricao", "Descrição resumida", "textarea"],
        ["prazo", "Prazo solicitado", "input"],
        ["prioridade", "Prioridade", "select", ["Alta", "Média", "Baixa", "Urgente"]],
        ["status", "Status", "input"]
      ],
      done: [
        ["processo", "Processo/demanda", "textarea"],
        ["atividade", "Atividade realizada", "textarea"],
        ["tempo", "Tempo aproximado", "input"],
        ["status", "Status", "input"],
        ["observacoes", "Observações", "textarea"]
      ]
    };

    function makeInput(type, value, options) {
      let el = document.createElement(type === "textarea" ? "textarea" : type === "select" ? "select" : "input");
      if (type === "select") {
        for (const opt of options || []) {
          const o = document.createElement("option");
          o.value = opt; o.textContent = opt; el.appendChild(o);
        }
      }
      el.value = value || "";
      return el;
    }

    function renderTable(name, rows) {
      const table = document.getElementById(name);
      const schema = schemas[name];
      table.innerHTML = "";
      const thead = document.createElement("thead");
      const h = document.createElement("tr");
      h.innerHTML = '<th class="col-num">Nº</th>' + schema.map(c => `<th>${c[1]}</th>`).join("") + '<th class="col-actions"></th>';
      thead.appendChild(h);
      table.appendChild(thead);
      const tbody = document.createElement("tbody");
      table.appendChild(tbody);
      (rows && rows.length ? rows : [{}]).forEach(row => addRow(name, row));
      renumber(name);
    }

    function addRow(name, data = {}) {
      const tbody = document.querySelector(`#${name} tbody`);
      const schema = schemas[name];
      const tr = document.createElement("tr");
      const num = document.createElement("td");
      num.className = "col-num rownum";
      tr.appendChild(num);
      for (const col of schema) {
        const td = document.createElement("td");
        const input = makeInput(col[2], data[col[0]], col[3]);
        input.dataset.field = col[0];
        td.appendChild(input);
        tr.appendChild(td);
      }
      const act = document.createElement("td");
      act.className = "col-actions";
      const del = document.createElement("button");
      del.type = "button";
      del.className = "small danger";
      del.textContent = "X";
      del.title = "Remover linha";
      del.onclick = () => { tr.remove(); renumber(name); };
      act.appendChild(del);
      tr.appendChild(act);
      tbody.appendChild(tr);
      renumber(name);
    }

    function renumber(name) {
      document.querySelectorAll(`#${name} .rownum`).forEach((td, i) => td.textContent = String(i + 1));
    }

    function collectTable(name) {
      const rows = [];
      document.querySelectorAll(`#${name} tbody tr`).forEach(tr => {
        const item = {};
        tr.querySelectorAll("[data-field]").forEach(el => item[el.dataset.field] = el.value.trim());
        if (Object.values(item).some(v => v)) rows.push(item);
      });
      return rows;
    }

    function val(id) { return document.getElementById(id).value.trim(); }
    function set(id, value) { document.getElementById(id).value = value || ""; }

    async function loadDefaults() {
      const response = await fetch("/api/defaults");
      const data = await response.json();
      for (const [key, value] of Object.entries(data.identity || {})) set(key, value);
      renderTable("planning", data.planning || []);
      renderTable("extra", data.extra || [{}]);
      renderTable("done", data.done || [{}]);
      for (const [key, value] of Object.entries(data.indicators || {})) {
        if (document.getElementById(key)) set(key, value);
      }
      set("observacoes", data.observacoes || "");
      set("validacao", "");
    }

    document.querySelectorAll("[data-add]").forEach(btn => {
      btn.addEventListener("click", () => addRow(btn.dataset.add));
    });

    document.getElementById("copyPlan").addEventListener("click", () => {
      const rows = collectTable("planning").map(item => ({
        processo: item.processo,
        atividade: item.descricao,
        tempo: "-",
        status: item.status,
        observacoes: ""
      }));
      renderTable("done", rows.length ? rows : [{}]);
    });

    document.getElementById("generate").addEventListener("click", async () => {
      const payload = {
        identity: {
          data: val("data"),
          colaborador: val("colaborador"),
          setor: val("setor"),
          funcao: val("funcao"),
          horario: val("horario"),
          responsavel: val("responsavel")
        },
        planning: collectTable("planning"),
        extra: collectTable("extra"),
        done: collectTable("done"),
        indicators: {
          produtividade: val("produtividade"),
          cumprimento: val("cumprimento"),
          urgente: val("urgente"),
          urgente_comentario: val("urgente_comentario"),
          dependencia: val("dependencia"),
          dependencia_comentario: val("dependencia_comentario"),
          retrabalho: val("retrabalho"),
          retrabalho_comentario: val("retrabalho_comentario"),
          sobrecarga: val("sobrecarga"),
          sobrecarga_comentario: val("sobrecarga_comentario"),
          acumulo: val("acumulo"),
          acumulo_comentario: val("acumulo_comentario")
        },
        observacoes: val("observacoes"),
        validacao: val("validacao")
      };
      const status = document.getElementById("resultText");
      const line = document.getElementById("result");
      status.textContent = "Gerando...";
      line.classList.add("show");
      const response = await fetch("/gerar", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify(payload)
      });
      const data = await response.json();
      if (!response.ok) {
        status.textContent = data.error || "Não foi possível gerar o relatório.";
        status.className = "status";
        return;
      }
      status.textContent = "Relatório gerado com sucesso.";
      status.className = "status ok";
      const link = document.getElementById("download");
      link.href = data.download_url;
      link.download = data.file_name;
      link.textContent = "Baixar " + data.file_name;
    });

    loadDefaults();
  </script>
</body>
</html>
"""


def text(value: object) -> str:
    return "" if value is None else str(value).strip()


def set_cell(cell, value: object) -> None:
    cell.text = text(value)


def rows_from_table(table, field_names, header_rows=1):
    rows = []
    for row in table.rows[header_rows:]:
        item = {}
        has_value = False
        for field, cell in zip(field_names, row.cells[1:]):
            value = " ".join(cell.text.split())
            item[field] = value
            has_value = has_value or bool(value)
        if has_value:
            rows.append(item)
    return rows


def read_defaults():
    doc = Document(str(TEMPLATE_PATH))
    today = datetime.now().strftime("%d/%m/%Y")
    identity = {
        "data": today,
        "colaborador": "Januária Medeiros",
        "setor": "Licitação",
        "funcao": "Analista de Compras",
        "horario": "07:30 às 16:30",
        "responsavel": "Stela Diniz",
    }
    try:
        table = doc.tables[0]
        identity.update({
            "colaborador": table.cell(0, 3).text.strip() or identity["colaborador"],
            "setor": table.cell(1, 1).text.strip() or identity["setor"],
            "funcao": table.cell(1, 3).text.strip() or identity["funcao"],
            "horario": table.cell(2, 1).text.strip() or identity["horario"],
            "responsavel": table.cell(2, 3).text.strip() or identity["responsavel"],
        })
    except Exception:
        pass

    planning = rows_from_table(doc.tables[1], ["processo", "descricao", "prioridade", "status"]) if len(doc.tables) > 1 else []
    indicators = {
        "produtividade": "Alto",
        "cumprimento": "Parcialmente",
        "urgente": "Sim",
        "dependencia": "Sim",
        "retrabalho": "Sim",
        "sobrecarga": "Sim",
        "acumulo": "Sim",
    }
    return {
        "identity": identity,
        "planning": planning or [{}],
        "extra": [{}],
        "done": [{}],
        "indicators": indicators,
        "observacoes": "As eventuais informações pontuais relativas às atividades foram registradas nos campos de observação específicos de cada tarefa.",
        "validacao": "",
    }


def clone_last_row(table) -> None:
    new_tr = copy.deepcopy(table.rows[-1]._tr)
    table._tbl.append(new_tr)


def set_body_row_count(table, desired: int, minimum: int = 1) -> None:
    desired = max(desired, minimum)
    while len(table.rows) - 1 < desired:
        clone_last_row(table)
    while len(table.rows) - 1 > desired:
        table._tbl.remove(table.rows[-1]._tr)


def fill_numbered_table(table, rows, fields, minimum=1):
    set_body_row_count(table, len(rows), minimum)
    for index, row in enumerate(table.rows[1:], start=1):
        data = rows[index - 1] if index - 1 < len(rows) else {}
        set_cell(row.cells[0], index)
        for offset, field in enumerate(fields, start=1):
            set_cell(row.cells[offset], data.get(field, ""))


def mark_two(selected: str) -> str:
    yes = selected.strip().lower().startswith("s")
    return f"({'x' if yes else ' '}) Sim ({' ' if yes else 'x'}) Não"


def mark_productivity(selected: str) -> str:
    selected = selected.strip().lower()
    return " ".join([
        f"({'x' if selected.startswith('alto') else ' '}) Alto",
        f"({'x' if selected.startswith('méd') or selected.startswith('med') else ' '}) Médio",
        f"({'x' if selected.startswith('baixo') else ' '}) Baixo",
    ])


def mark_cumprimento(selected: str) -> str:
    selected = selected.strip().lower()
    return " ".join([
        f"({'x' if selected.startswith('sim') else ' '}) Sim",
        f"({'x' if selected.startswith('parc') else ' '}) Parcialmente",
        f"({'x' if selected.startswith('não') or selected.startswith('nao') else ' '}) Não",
    ])


def fill_indicators(table, indicators):
    rows = [
        ("Nível de produtividade percebido", mark_productivity(indicators.get("produtividade", "Alto")), ""),
        ("Cumprimento das atividades planejadas", mark_cumprimento(indicators.get("cumprimento", "Parcialmente")), ""),
        ("Houve demanda urgente não planejada?", mark_two(indicators.get("urgente", "Sim")), indicators.get("urgente_comentario", "")),
        ("Houve dependência de outro setor? Qual?", mark_two(indicators.get("dependencia", "Sim")), indicators.get("dependencia_comentario", "")),
        ("Houve retrabalho? Por quê?", mark_two(indicators.get("retrabalho", "Sim")), indicators.get("retrabalho_comentario", "")),
        ("Houve sobrecarga demandas? Discorra.", mark_two(indicators.get("sobrecarga", "Sim")), indicators.get("sobrecarga_comentario", "")),
        ("Houve acúmulo de demandas? Discorra.", mark_two(indicators.get("acumulo", "Sim")), indicators.get("acumulo_comentario", "")),
    ]
    set_body_row_count(table, len(rows), len(rows))
    for row, values in zip(table.rows[1:], rows):
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)


def fill_single_column_table(table, value: str, minimum=4):
    lines = [line.strip() for line in text(value).splitlines() if line.strip()]
    desired = max(len(lines), minimum)
    while len(table.rows) < desired:
        clone_last_row(table)
    while len(table.rows) > desired:
        table._tbl.remove(table.rows[-1]._tr)
    for index, row in enumerate(table.rows):
        set_cell(row.cells[0], lines[index] if index < len(lines) else "")


def safe_filename(period: str) -> str:
    cleaned = text(period).upper()
    cleaned = cleaned.replace("/", "_").replace("\\", "_")
    cleaned = re.sub(r"[^A-Z0-9_ AÀ-Ú.-]+", "", cleaned)
    cleaned = re.sub(r"\s+", "_", cleaned).strip("_")
    return f"RELATORIO_DE_PRODUTIVIDADE_{cleaned or datetime.now().strftime('%d_%m_%Y')}.docx"


def generate_report(payload):
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc = Document(str(TEMPLATE_PATH))
    identity = payload.get("identity", {})

    t0 = doc.tables[0]
    set_cell(t0.cell(0, 1), identity.get("data", ""))
    set_cell(t0.cell(0, 3), identity.get("colaborador", ""))
    set_cell(t0.cell(1, 1), identity.get("setor", ""))
    set_cell(t0.cell(1, 3), identity.get("funcao", ""))
    set_cell(t0.cell(2, 1), identity.get("horario", ""))
    set_cell(t0.cell(2, 3), identity.get("responsavel", ""))

    fill_numbered_table(doc.tables[1], payload.get("planning", []), ["processo", "descricao", "prioridade", "status"], minimum=5)
    fill_numbered_table(doc.tables[2], payload.get("extra", []), ["processo", "descricao", "prazo", "prioridade", "status"], minimum=5)
    fill_numbered_table(doc.tables[3], payload.get("done", []), ["processo", "atividade", "tempo", "status", "observacoes"], minimum=5)
    fill_indicators(doc.tables[4], payload.get("indicators", {}))
    fill_single_column_table(doc.tables[5], payload.get("observacoes", ""), minimum=4)
    fill_single_column_table(doc.tables[6], payload.get("validacao", ""), minimum=4)

    name = safe_filename(identity.get("data", ""))
    target = OUTPUT_DIR / name
    counter = 2
    while target.exists():
        target = OUTPUT_DIR / f"{target.stem}_{counter}.docx"
        counter += 1
    doc.save(str(target))
    return target


class Handler(BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        return

    def send_json(self, data, status=200):
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self):
        parsed = urlparse(self.path)
        if parsed.path == "/":
            body = HTML.encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)
            return
        if parsed.path == "/api/defaults":
            self.send_json(read_defaults())
            return
        if parsed.path == "/download":
            params = parse_qs(parsed.query)
            name = Path(params.get("file", [""])[0]).name
            target = OUTPUT_DIR / name
            if not target.exists():
                self.send_json({"error": "Arquivo não encontrado."}, 404)
                return
            body = target.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", f'attachment; filename="{target.name}"')
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)
            return
        self.send_json({"error": "Página não encontrada."}, 404)

    def do_POST(self):
        if self.path != "/gerar":
            self.send_json({"error": "Página não encontrada."}, 404)
            return
        try:
            length = int(self.headers.get("Content-Length", "0"))
            payload = json.loads(self.rfile.read(length).decode("utf-8"))
            target = generate_report(payload)
            self.send_json({
                "file_name": target.name,
                "download_url": f"/download?file={target.name}",
            })
        except Exception as exc:
            self.send_json({"error": f"Erro ao gerar relatório: {exc}"}, 500)


def find_port(start=DEFAULT_PORT):
    for port in range(start, start + 50):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            try:
                s.bind((HOST, port))
                return port
            except OSError:
                continue
    raise RuntimeError("Não encontrei uma porta livre para iniciar o gerador.")


def local_browser_url(port: int) -> str:
    host = "127.0.0.1" if HOST in ("0.0.0.0", "") else HOST
    return f"http://{host}:{port}/"


def lan_ip() -> str:
    try:
        with socket.socket(socket.AF_INET, socket.SOCK_DGRAM) as s:
            s.connect(("8.8.8.8", 80))
            return s.getsockname()[0]
    except OSError:
        return socket.gethostbyname(socket.gethostname())


def main():
    if not TEMPLATE_PATH.exists():
        print(f"Modelo não encontrado: {TEMPLATE_PATH}")
        sys.exit(1)
    OUTPUT_DIR.mkdir(exist_ok=True)
    port = find_port()
    server = ThreadingHTTPServer((HOST, port), Handler)
    url = local_browser_url(port)
    print(f"Gerador aberto neste computador: {url}")
    if HOST == "0.0.0.0":
        print("")
        print("Link para outro computador na mesma rede:")
        print(f"http://{lan_ip()}:{port}/")
        print("")
        print("Deixe esta janela aberta enquanto a outra pessoa usa o link.")
    threading.Timer(0.7, lambda: webbrowser.open(url)).start()
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nGerador encerrado.")


if __name__ == "__main__":
    main()
